from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Auralis_Risk_Solution_Walkthrough.docx"
MODEL = json.loads((ROOT / "models" / "auralis-xgb-210k-v1.json").read_text(encoding="utf-8"))
HOLDOUT = json.loads((ROOT / "models" / "auralis-xgb-210k-holdout.json").read_text(encoding="utf-8"))
BASELINE = json.loads((ROOT / "models" / "fraudguard-linear-v1.json").read_text(encoding="utf-8"))

NAVY = "0B2545"
BLUE = "2E74B5"
CYAN = "1597A5"
DARK = "17212B"
MUTED = "667085"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "7A5A00"
RED = "9B1C1C"
PAGE_WIDTH_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, name="Calibri", size=None, color=DARK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "1")
    tr_pr.append(header)


def set_table_borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_custom_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    bullet_abs_id = abstract_id + 1
    num_id = max(existing_num, default=0) + 1
    bullet_num_id = num_id + 1

    def create_abstract(identifier, fmt, text_value):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(identifier))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        level.append(lvl_text)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        level.append(p_pr)
        # OOXML requires all abstract numbering definitions to precede <w:num>.
        first_num = numbering.find(qn("w:num"))
        if first_num is None:
            numbering.append(abstract)
        else:
            numbering.insert(list(numbering).index(first_num), abstract)

    def create_num(identifier, abstract_identifier):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(identifier))
        abstract_id_node = OxmlElement("w:abstractNumId")
        abstract_id_node.set(qn("w:val"), str(abstract_identifier))
        num.append(abstract_id_node)
        numbering.append(num)

    create_abstract(abstract_id, "decimal", "%1.")
    create_abstract(bullet_abs_id, "bullet", "•")
    create_num(num_id, abstract_id)
    create_num(bullet_num_id, bullet_abs_id)
    return num_id, bullet_num_id


def fresh_num_id(doc, source_num_id):
    """Create a new numbering instance so every list starts independently."""
    numbering = doc.part.numbering_part.element
    source = next(
        node for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) == str(source_num_id)
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    identifier = max(existing, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(identifier))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    numbering.append(num)
    return identifier


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_list(doc, items, ordered=False):
    # Manual markers render consistently in Word and LibreOffice and avoid
    # cross-list continuation caused by application-specific numbering state.
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.167
        marker = f"{index}. " if ordered else "\u2022 "
        run = paragraph.add_run(marker)
        set_font(run)
        if isinstance(item, tuple):
            label, detail = item
            run = paragraph.add_run(label)
            set_font(run, bold=True)
            run = paragraph.add_run(detail)
            set_font(run)
        else:
            run = paragraph.add_run(item)
            set_font(run)


def add_callout(doc, label, text, fill=PALE, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    mark_repeat_header(table.rows[0])
    set_table_geometry(table, [PAGE_WIDTH_DXA], indent=120)
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    set_font(run, size=10.5, color=accent, bold=True)
    run = paragraph.add_run(text)
    set_font(run, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    table.rows[0].cells
    mark_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_font(run, size=9, color=NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            run = paragraph.add_run(str(value))
            set_font(run, size=font_size)
    set_table_geometry(table, widths)
    return table


def add_section_title(doc, number, title, lead):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"STEP {number:02d}")
    set_font(run, size=9.5, color=CYAN, bold=True)
    doc.add_heading(title, level=1)
    paragraph = doc.add_paragraph(lead)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.runs[0]
    set_font(run, size=11.5, color=MUTED, italic=True)


def add_page_break(doc, even_page=False):
    # Attach the break to the final paragraph so a nearly full page cannot push
    # an empty break paragraph onto the next page and create a blank sheet.
    paragraph = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    configure_styles(doc)
    # Explicit odd/even header parts keep LibreOffice and Word rendering aligned.
    doc.settings.odd_and_even_pages_header_footer = True
    decimal_id, bullet_id = add_custom_numbering(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("AURALIS RISK  |  MASTERCARD INNOVATION CHALLENGE 2026")
    set_font(run, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    even_header = section.even_page_header
    paragraph = even_header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("AURALIS RISK  |  MASTERCARD INNOVATION CHALLENGE 2026")
    set_font(run, size=8.5, color=MUTED, bold=True)
    add_page_number(section.even_page_footer.paragraphs[0])

    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AI DEFENSE LAB FOR PAYMENT SECURITY")
    set_font(run, size=10, color=CYAN, bold=True)
    kicker.paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Auralis Risk")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Identify. Generate. Defend. Learn.")
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(34)
    run = descriptor.add_run("Solution Walkthrough | Mastercard Innovation Challenge @ GFF 2026")
    set_font(run, size=11, color=MUTED)

    add_callout(
        doc,
        "One-line pitch",
        "Auralis Risk turns emerging fraud into a governed synthetic moving target so defenders can discover, test, detect, and learn before an attack reaches the payment decision path.",
        fill="EAF5F7",
        accent=CYAN,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(22)
    meta = add_table(
        doc,
        ["Submission artifact", "Current prototype"],
        [
            ("System shape", "Modular monolith with offline training and deterministic real-time scoring"),
            ("Dataset", "210,000 traceable synthetic training transactions plus one unseen attack-family holdout"),
            ("Model", f"Portable 280-tree XGBoost + transparent rules ({MODEL['model_version']})"),
            ("Safety", "Synthetic-only data; human-approved GenAI scenarios; no LLM in the authorization path"),
        ],
        [2300, 7060],
        font_size=9.5,
    )

    add_page_break(doc)
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph(
        "Generative AI makes payment fraud cheaper to vary, faster to scale, and harder for static defenses to anticipate. "
        "Auralis Risk answers the competition brief with one measurable closed loop: discover plausible attack families, "
        "generate safe synthetic payment behavior, score it through a live-style defense, identify misses and false positives, "
        "and route those gaps into reviewed hardening work."
    )
    add_callout(
        doc,
        "Competition fit",
        "The challenge provides no dataset. Creating a high-fidelity, versioned simulation dataset is therefore a core deliverable, not a workaround.",
    )
    doc.add_heading("What the judges can see", level=2)
    add_list(
        doc,
        [
            ("Attack diversity - ", "22 structured families spanning identity, device, merchant, network, bot, and payment behavior."),
            ("Simulation fidelity - ", "seeded, constrained transactions with provenance, time order, fictional entities, and explicit synthetic labels."),
            ("Detection efficacy - ", "precision, recall, F1, false-positive rate, confusion matrix, reason codes, and latency."),
            ("Novelty - ", "a completely excluded attack-family holdout and a linear-baseline comparison create measurable evidence of generalization."),
            ("Live-payment feasibility - ", "a dependency-light synchronous API keeps generative AI outside the authorization path."),
        ],
        False,
    )
    doc.add_heading("Architecture decision", level=2)
    doc.add_paragraph(
        "The planning brief proposed six services. For the hackathon, the same boundaries are implemented as modules inside one deployable process. "
        "This makes the five-minute demo reliable while preserving clean extraction points for a Python model service, Redis, PostgreSQL, and event streaming."
    )
    add_table(
        doc,
        ["Path", "Components", "Purpose"],
        [
            ("Offline", "Scenario registry, generator, training", "Discover and harden"),
            ("Nearline", "Evaluation, feedback, drift candidates", "Measure and learn"),
            ("Real-time", "Features, model adapter, rules, policy", "Score and decide"),
        ],
        [1500, 3500, 4360],
    )

    add_section_title(
        doc,
        1,
        "Identify: build a defensible attack landscape",
        "The red team is a research and simulation layer, not a mechanism for interacting with victims or live payment rails.",
    )
    attacks = [
        ("ATO_001", "Account takeover", "New device, velocity, location shift"),
        ("CNP_001", "Card-not-present burst", "Remote payment, device novelty, amount deviation"),
        ("MULE_001", "Mule network", "Shared entity, new payee, velocity"),
        ("BOT_001", "Bot card testing", "Small bursts, repeated attempts, velocity"),
        ("REFUND_001", "Refund abuse", "Merchant risk, clustering, amount shift"),
        ("UPI_001", "Urgent instant-payment scam", "New payee, unusual time, amount shift"),
        ("SYNID_001", "Synthetic identity", "New account, inconsistent identity, velocity"),
        ("LAUNDER_001", "Transaction layering", "Split/merge behavior, graph density, new payee"),
        ("PROMO_001", "Promotion abuse", "Shared device, new accounts, small bursts"),
        ("FRIENDLY_001", "Friendly fraud", "Stable device, amount anomaly, merchant context"),
        ("SIMSWAP_001", "Recovery-channel takeover", "New device, identity mismatch, location shift"),
        ("TOKEN_001", "Token provisioning abuse", "Device novelty, shared infrastructure, remote use"),
        ("QR_001", "QR destination substitution", "New payee, merchant mismatch, geography"),
        ("BNPL_001", "Cross-provider credit bust-out", "New account, velocity, high-value deviation"),
        ("INVOICE_001", "Business payment redirection", "New beneficiary, amount shift, merchant context"),
        ("LOYALTY_001", "Loyalty value theft", "New device, redemption velocity, small-value bursts"),
        ("SUBSCRIPTION_001", "Synthetic subscription farm", "Shared device, remote payment, new accounts"),
        ("MERCHANT_001", "Transaction laundering", "Merchant risk, graph density, amount distribution"),
        ("NFC_001", "Contactless proximity relay", "Location contradiction, stable device, amount shift"),
        ("REMIT_001", "Remittance corridor abuse", "Cross-border paths, new payees, graph density"),
        ("PAYROLL_001", "Payroll redirection", "New destination, amount shift, identity mismatch"),
        ("GIFT_001", "Gift-card conversion cascade", "Small-value bursts, velocity, merchant risk"),
    ]
    add_table(doc, ["ID", "Attack family", "Observable defensive signals"], attacks, [1400, 3000, 4960], font_size=9)
    doc.add_heading("Scenario governance", level=2)
    add_list(
        doc,
        [
            "Every scenario has a stable ID, version, severity, feature coverage, and simulation-ready state.",
            "Mutation is limited to safe attributes such as timing, amount distribution, device novelty, geography, merchant context, and relationship patterns.",
            "Every generated dataset records its scenario version and seed so results can be reproduced.",
            "Human review is required before a proposed scenario enters the approved registry.",
        ],
        False,
    )
    add_callout(
        doc,
        "Responsible-use boundary",
        "No real card numbers, CVVs, credentials, customer PII, phishing delivery, or unauthorized transaction execution is used or produced.",
        fill="FBEFEF",
        accent=RED,
    )
    doc.add_heading("Governed GenAI threat discovery", level=2)
    add_list(
        doc,
        [
            "An optional Groq-hosted analyst proposes strictly structured, defensive simulation drafts; a labeled local analyst fallback keeps the demo reproducible without an API key.",
            "Every response is schema-validated, safety-filtered, bounded to synthetic parameters, and stored with provider/model provenance.",
            "A new draft starts as PENDING_REVIEW. Simulation remains disabled until a human explicitly approves it; rejected drafts cannot run.",
            "False negatives can seed new reviewed drafts, but neither GenAI nor feedback can train or promote a production model automatically.",
        ],
        False,
    )
    genai_image = ROOT / "docs" / "assets" / "auralis-genai-review.png"
    genai_paragraph = doc.add_paragraph()
    genai_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    genai_shape = genai_paragraph.add_run().add_picture(str(genai_image), width=Inches(6.25))
    genai_shape._inline.docPr.set("name", "Auralis Risk governed GenAI threat review")
    genai_shape._inline.docPr.set("descr", "Threat analyst draft marked pending review, with approval, rejection, and disabled simulation controls.")
    genai_caption = doc.add_paragraph()
    genai_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    genai_caption.paragraph_format.space_after = Pt(0)
    set_font(genai_caption.add_run("Figure 1. A threat draft cannot generate evidence before human approval."), size=9, color=MUTED, italic=True)

    add_page_break(doc)
    add_section_title(
        doc,
        2,
        "Generate: create a traceable synthetic dataset",
        "Because the challenge supplies no dataset, simulation quality is part of the product and must be audited like a model.",
    )
    doc.add_heading("Current dataset", level=2)
    add_table(
        doc,
        ["Property", "Implementation"],
        [
            ("Size", "210,000 training/validation/test feature rows"),
            ("Coverage", "21 training attack families; LAUNDER_001 excluded as holdout"),
            ("Class balance", "Approximately 25% fraud per scenario for the current controlled benchmark"),
            ("Reproducibility", "Fixed seeds, generator version, feature version, and scenario provenance"),
            ("Privacy", "All identifiers and events are fictional; rows are marked synthetic"),
            ("Storage", "The versioned 210,000-row release archive, manifest, model artifact, and evaluation evidence are committed"),
        ],
        [2200, 7160],
    )
    doc.add_heading("Generation sequence", level=2)
    add_list(
        doc,
        [
            "Select an approved scenario and fixed seed.",
            "Generate baseline legitimate behavior across customers, merchants, devices, amounts, channels, and time.",
            "Condition fraud rows on the scenario's observable signals.",
            "Process transactions in timestamp order through the production FeatureEngine.",
            "Hash customer/scenario entities into leakage-safe train, validation, and test partitions.",
            "Write labels, provenance, feature version, and holdout metadata.",
        ],
        True,
    )
    doc.add_heading("Fidelity controls", level=2)
    add_list(
        doc,
        [
            ("Implemented - ", "12 automated gates covering deterministic replay, schema, class rates, hard-negative coverage, amount overlap, correlation, feature ranges, and privacy."),
            ("Next - ", "comparison with an authorized aggregate reference distribution, sequence similarity, and graph-structure reports."),
        ],
        False,
    )
    add_callout(
        doc,
        "Honest limitation",
        "The fidelity suite validates generator consistency and controlled class overlap. Without an authorized production reference, it cannot prove production realism.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_page_break(doc)
    add_section_title(
        doc,
        3,
        "Feature engineering: represent behavior consistently",
        "Training and inference call the same stateful feature module to reduce training-serving skew.",
    )
    features = [
        ("velocity_1h", "Transactions observed for the customer in the prior hour", "Velocity"),
        ("amount_deviation", "Distance from the customer's prior average amount", "Behavior"),
        ("new_device", "Device not observed in prior customer history", "Device"),
        ("shared_device_count", "Customers associated with the device", "Graph proxy"),
        ("location_shift", "Country differs from current home-country assumption", "Context"),
        ("new_payee", "Payment introduces a new recipient", "Payment"),
        ("card_not_present", "Remote/e-commerce transaction", "Channel"),
        ("unusual_hour", "Event falls in a configured high-risk UTC window", "Time"),
        ("new_account", "Account age is below 30 days", "Account"),
        ("identity_mismatch", "Synthetic identity attributes are inconsistent", "Identity"),
        ("merchant_risk", "Normalized merchant-risk context", "Merchant"),
    ]
    add_table(doc, ["Feature", "Meaning", "Domain"], features, [2400, 5260, 1700], font_size=8.7)
    doc.add_heading("Parity rules", level=2)
    add_list(
        doc,
        [
            "Process events in timestamp order.",
            "Version the feature schema independently from the model.",
            "Log the feature version with every decision.",
            "Treat stateful history and device relationships as part of the scoring contract.",
        ],
        False,
    )

    add_page_break(doc)
    add_section_title(
        doc,
        4,
        "Defend: train an explainable model and safe policy",
        "A compact nonlinear detector improves recall while remaining inexpensive and reproducible to serve.",
    )
    doc.add_heading("Why portable XGBoost", level=2)
    add_list(
        doc,
        [
            "Gradient-boosted trees capture nonlinear interactions between velocity, deviation, device, graph, identity, and merchant signals.",
            "A portable JSON artifact supports dependency-light inference in the Node service without a separate Python runtime.",
            "Deterministic training on the locked 210,000-row release makes the comparison against the original linear baseline reproducible.",
            "If the artifact is missing or invalid, the system continues with transparent rule scoring.",
        ],
        False,
    )
    doc.add_heading("Training flow", level=2)
    add_list(
        doc,
        [
            "Load versioned feature rows.",
            "Preserve customer-entity train, validation, and test partitions to reduce leakage.",
            "Train a 280-tree XGBoost classifier with fixed seed and controlled class weighting.",
            "Select a validation threshold under a false-positive constraint.",
            "Evaluate the untouched test split and the completely excluded LAUNDER_001 family, then save locked artifacts.",
        ],
        True,
    )
    add_table(
        doc,
        ["Artifact field", "Value"],
        [
            ("Model version", MODEL["model_version"]),
            ("Model type", MODEL["model_type"]),
            ("Feature version", MODEL["feature_version"]),
            ("Validation threshold", MODEL["decision_threshold"]),
            ("Training examples", MODEL["training_manifest"]["examples"]["train"]),
            ("Validation examples", MODEL["training_manifest"]["examples"]["validation"]),
            ("Test examples", MODEL["training_manifest"]["examples"]["test"]),
        ],
        [2900, 6460],
    )
    doc.add_heading("Real-time ensemble and policy", level=2)
    doc.add_paragraph(
        "The runtime score blends 70% trained fraud probability with 30% transparent rule score. "
        "Configurable thresholds map the result to ALLOW, STEP_UP, REVIEW, or BLOCK. "
        "Every response records probability, rule score, reason codes, feature/model versions, scoring mode, and latency."
    )

    add_section_title(
        doc,
        5,
        "Evaluate: measure security and customer friction",
        "F1 and recall matter, but false positives, explainability, and latency are equally visible in the prototype.",
    )
    test_metrics = MODEL["metrics"]["test"]
    holdout_metrics = HOLDOUT["metrics"]
    baseline_metrics = BASELINE["metrics"]["test"]
    add_table(
        doc,
        ["Evaluation", "Precision", "Recall", "F1", "False-positive rate"],
        [
            ("Entity-aware test split", f"{test_metrics['precision']:.3f}", f"{test_metrics['recall']:.3f}", f"{test_metrics['f1']:.3f}", f"{test_metrics['false_positive_rate']:.3f}"),
            ("Excluded LAUNDER_001 holdout", f"{holdout_metrics['precision']:.3f}", f"{holdout_metrics['recall']:.3f}", f"{holdout_metrics['f1']:.3f}", f"{holdout_metrics['false_positive_rate']:.3f}"),
            ("Original linear test baseline", f"{baseline_metrics['precision']:.3f}", f"{baseline_metrics['recall']:.3f}", f"{baseline_metrics['f1']:.3f}", f"{baseline_metrics['false_positive_rate']:.3f}"),
        ],
        [3000, 1500, 1500, 1400, 1960],
        font_size=9,
    )
    doc.add_heading("How to read the result", level=2)
    add_list(
        doc,
        [
            ("Test split - ", "measures held-out customers within the 21 known attack families."),
            ("Novel holdout - ", "uses LAUNDER_001, which is absent from the training dataset."),
            ("Holdout - ", "measures 10,000 LAUNDER_001 rows from a family completely absent from training."),
            ("Measured lift - ", "XGBoost improves test F1 by 3.42 points and recall by 5.70 points over the linear baseline with essentially unchanged false-positive rate."),
        ],
        False,
    )
    add_callout(
        doc,
        "Metric caution",
        "Hard negatives reduce synthetic separability and expose the security/customer-friction tradeoff. These values remain prototype evidence, not production claims.",
        fill="FFF8E8",
        accent=GOLD,
    )
    doc.add_heading("Prototype evidence", level=2)
    image_path = ROOT / "docs" / "assets" / "auralis-xgboost-console.png"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.25))
    doc_pr = shape._inline.docPr
    doc_pr.set("name", "Auralis Risk holdout evaluation dashboard")
    doc_pr.set("descr", "Live feature-vector form showing a high-risk classification, fraud probability, risk score, model version, and inference latency.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)
    run = caption.add_run("Figure 2. Browser-verified local XGBoost inference with explicit artifact version and latency.")
    set_font(run, size=9, color=MUTED, italic=True)

    add_page_break(doc)
    add_section_title(
        doc,
        6,
        "Learn: turn gaps into reviewed hardening work",
        "The feedback loop creates candidates for improvement; it never silently updates the live model.",
    )
    add_table(
        doc,
        ["Input", "Classification", "Controlled next action"],
        [
            ("Predicted FRAUD, actual LEGITIMATE", "False positive", "Create benign hard negative; review threshold impact"),
            ("Predicted LEGITIMATE, actual FRAUD", "False negative", "Seed safe defense-guided mutation"),
            ("Prediction matches outcome", "Correct", "Retain for monitoring and calibration"),
        ],
        [2900, 2100, 4360],
        font_size=9,
    )
    doc.add_heading("Defense-guided mutation", level=2)
    add_list(
        doc,
        [
            "Replay the unseen scenario and select actual false negatives.",
            "Create traceable amount, payee, device, and merchant-context stress variants.",
            "Replay each variant through exactly the same feature and decision path.",
            "Mark the batch HUMAN_REVIEW_REQUIRED; never change the active artifact.",
            "Next, train an isolated candidate only from approved batches.",
            "Compare locked metrics before any human-approved promotion.",
        ],
        True,
    )
    add_callout(
        doc,
        "Novelty mechanism",
        "The red team does not generate random fraud. It searches the defender's uncertainty and blind spots, then creates traceable, plausible synthetic variants for stress testing.",
        fill="EAF5F7",
        accent=CYAN,
    )

    add_page_break(doc)
    doc.add_heading("Real-world Feasibility", level=1)
    add_table(
        doc,
        ["Stage", "Deployment mode", "Control objective"],
        [
            ("1", "Offline simulation", "Validate diversity, fidelity, model lift, and limitations"),
            ("2", "Shadow mode", "Score authorized live-like traffic without changing approvals"),
            ("3", "Step-up recommendation", "Use medium risk for extra authentication or review"),
            ("4", "Selective blocking", "Block only calibrated high-confidence patterns with rollback"),
            ("5", "Continuous hardening", "Schedule reviewed mutations, retraining, and governance"),
        ],
        [1000, 2700, 5660],
        font_size=9,
    )
    doc.add_heading("Why the prototype can evolve", level=2)
    add_list(
        doc,
        [
            "Module boundaries allow feature, model, generator, and monitoring components to split into independent services when scale requires it.",
            "The model adapter now serves the portable XGBoost artifact locally and can later call a separately governed model service without changing the score API.",
            "Redis can hold online feature state, PostgreSQL can persist scenarios and feedback, and a queue can carry nearline events.",
            "The hot path remains deterministic, versioned, observable, and independent of generative-model availability.",
        ],
        False,
    )
    doc.add_heading("Responsible-use and security controls", level=2)
    add_list(
        doc,
        [
            "Synthetic and authorized test data only.",
            "No real payment credentials or customer PII in the red-team layer.",
            "Every scenario and decision carries provenance and version metadata.",
            "No uncontrolled self-training or automatic model promotion.",
            "Prototype metrics are never presented as production readiness.",
        ],
        False,
    )

    add_page_break(doc)
    doc.add_heading("Five-minute Demonstration", level=1)
    add_list(
        doc,
        [
            ("0:00-0:35 - Problem. ", "Explain why GenAI changes the speed and variety of payment fraud."),
            ("0:35-1:20 - Identify. ", "Generate a GenAI threat draft, show the blocked simulation button, approve it, then run the safe synthetic scenario."),
            ("1:20-2:00 - Generate. ", "Launch a seeded replay and explain provenance, 210,000-row lineage, and fidelity controls."),
            ("2:00-2:55 - Defend. ", "Enter a feature vector in the website and show the live XGBoost probability, decision, model version, and latency."),
            ("2:50-3:40 - Learn. ", "Generate a governed mutation batch from observed holdout misses."),
            ("3:40-4:30 - Feasibility. ", "Explain deterministic hot path, fallback, staged deployment, and human approval."),
            ("4:30-5:00 - Close. ", "State limitations and the defense-guided mutation roadmap."),
        ],
        True,
    )
    doc.add_heading("Implementation roadmap", level=2)
    add_table(
        doc,
        ["Priority", "Next deliverable", "Evidence of completion"],
        [
            ("P0", "Candidate training", "Before/after artifact on approved mutation batches"),
            ("P0", "Reference fidelity", "Compare against authorized aggregate distributions when available"),
            ("P1", "Calibration benchmark", "Probability calibration and cost-sensitive threshold comparison on locked data"),
            ("P1", "Graph intelligence", "Shared-entity and ring-density lift"),
            ("P1", "Observability", "Drift test, metrics endpoint, durable audit trail"),
        ],
        [1200, 3300, 4860],
        font_size=9,
    )
    doc.add_heading("Submission checklist", level=2)
    add_list(
        doc,
        [
            "Runnable repository covers Identify, Generate, Defend, and Learn.",
            "DOCX explains attacks, generation, detection efficacy, and live-payment feasibility.",
            "Web prototype demonstrates the closed loop with a presentable UI.",
            "All synthetic metrics and limitations are labeled honestly.",
            "Public deployment and links are tested without requiring judge login.",
        ],
        False,
    )
    doc.add_heading("Sources", level=2)
    sources = [
        "Mastercard Innovation Challenge 2026 private Kaggle overview (accessed 22 August 2026): https://www.kaggle.com/competitions/mastercard-innovation-challenge-2026",
        "Mastercard Innovation Challenge @ GFF 2026 public Luma brief (accessed 22 August 2026): https://luma.com/kyz978xv",
        "Auralis Risk repository implementation, portable XGBoost artifact, automated tests, and browser-verified prototype.",
        "User-provided planning documents: team work plan and Mastercard GFF 2026 AI Defense Lab Project Report.",
    ]
    add_list(doc, sources, False)

    doc.core_properties.title = "Auralis Risk Solution Walkthrough"
    doc.core_properties.subject = "Mastercard Innovation Challenge 2026 - AI Defense Lab for Payment Security"
    doc.core_properties.author = "Auralis Risk Team"
    doc.core_properties.keywords = "fraud detection, synthetic data, red team, blue team, payment security"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
